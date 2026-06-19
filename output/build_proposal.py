#!/usr/bin/env python3
"""Build the Sotatek/Alba vendor proposal .docx from the company template.

Strategy: load the template (keeps the green header bar + named Title/Heading
styles), strip the example body, then inject content derived from the
production-readiness gap-analysis plan. Body font forced to Times New Roman
per the template spec. Tables are real Word tables with manual borders.
"""
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = "Sotatek_project_proposal_v1.0_template.docx"
OUTPUT = "Sotatek_Alba_Proposal_v1.1.docx"

FONT = "Times New Roman"
GREEN = RGBColor(0x00, 0xAB, 0x44)

doc = Document(TEMPLATE)

# ---------------------------------------------------------------- body reset
body = doc.element.body
sectPr = body.find(qn("w:sectPr"))  # keep the section props (header refs)
for child in list(body):
    if child is not sectPr:
        body.remove(child)

# ensure first-page (green-bar) header renders
if sectPr is not None and sectPr.find(qn("w:titlePg")) is None:
    sectPr.append(OxmlElement("w:titlePg"))

# ---------------------------------------------------------- font: Times NR
def force_font(style_name):
    try:
        st = doc.styles[style_name]
    except KeyError:
        return
    st.font.name = FONT
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), FONT)

for s in ("Normal", "Title", "Subtitle", "Heading1", "Heading2",
          "Heading3", "Heading4"):
    force_font(s)
doc.styles["Normal"].font.size = Pt(11)

# --------------------------------------------------------------- helpers
def _style(name, fallback):
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback

H1 = _style("Heading1", "Heading 1")
H2 = _style("Heading2", "Heading 2")
H3 = _style("Heading3", "Heading 3")

def heading(text, level):
    p = doc.add_paragraph(text, style={1: H1, 2: H2, 3: H3}[level])
    return p

def para(text="", bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = FONT
        if color:
            r.font.color.rgb = color
    return p

def bullet(text, bold_lead=None, level=0):
    p = doc.add_paragraph(style="Normal")
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3 + 0.3 * level)
    pf.space_after = Pt(2)
    r0 = p.add_run("•  ")
    r0.font.name = FONT
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.bold = True
        rb.font.name = FONT
        rb.font.size = Pt(11)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    return p

def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tcPr.append(borders)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def add_table(headers, rows, widths=None, section_rows=None, font_size=9.5):
    """section_rows: set of row-indices (in `rows`) that are full-width
    section header bands."""
    section_rows = section_rows or set()
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = FONT
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(c, "00AB44")
        _set_cell_border(c)
    # body
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        if i in section_rows:
            # merge full width as a band
            merged = cells[0]
            for k in range(1, len(headers)):
                merged = merged.merge(cells[k])
            merged.text = ""
            p = merged.paragraphs[0]
            run = p.add_run(str(row[0]))
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = FONT
            run.font.color.rgb = GREEN
            _shade(merged, "E8F5EC")
            _set_cell_border(merged)
            continue
        for j, val in enumerate(row):
            c = cells[j]
            c.text = ""
            p = c.paragraphs[0]
            run = p.add_run("" if val is None else str(val))
            run.font.size = Pt(font_size)
            run.font.name = FONT
            if j == 0 or (headers and headers[j].upper() == "TOTAL"):
                pass
            # bold TOTAL rows
            if str(row[0]).strip().upper().startswith("TOTAL"):
                run.bold = True
            _set_cell_border(c)
    if widths:
        for col, w in zip(t.columns, widths):
            for c in col.cells:
                c.width = Inches(w)
    return t

# ============================================================ COVER
para("SotaTek — Software Development & Consulting   |   Vietnam", bold=True,
     size=11, color=GREEN)
title = doc.add_paragraph(style="Title")
tr = title.add_run("Alba — Golf Social Platform")
tr.font.name = FONT
sub = doc.add_paragraph(style=_style("Subtitle", "Normal"))
sr = sub.add_run("Production-Readiness Hardening & Go-Live Remediation — Technical Proposal & Estimate")
sr.italic = True
sr.font.name = FONT
sr.font.size = Pt(13)
para("19th June 2026  ·  Version 1.1", size=11)
para("Prepared for: Alba (Client)   ·   Prepared by: SotaTek", size=11)

# ============================================================ OVERVIEW
heading("OVERVIEW", 1)
para("Alba is a mobile-first golf social platform (iOS + Android) that lets "
     "golfers post and discover games/rounds, match with players who fit their "
     "preferences, coordinate via real-time chat and notifications, and split "
     "the cost of play through in-app payments. The product is already built — "
     "a NestJS/Prisma/PostgreSQL backend and an Expo/React Native mobile app, "
     "both authenticating the same user through Firebase. This engagement takes "
     "that existing codebase from its current ~70–75% state to a secure, "
     "reliable, production-grade go-live.")

heading("Core Business Model", 3)
bullet("two-sided social marketplace for golf games and rounds, mobile-first.",
       bold_lead="Social Sports Platform: ")
bullet("post a game with requirements → players request to join → organizer "
       "approves → cost split and collected in-app.", bold_lead="Game Matching & Coordination: ")
bullet("Stripe Connect collects player payments, holds funds ~2 days for "
       "disputes, then pays out organizers.", bold_lead="Marketplace Payments (Stripe Connect): ")
bullet("chat, live notifications, leaderboards, posts, groups and "
       "relationships drive retention.", bold_lead="Engagement & Community: ")
bullet("course catalogue, complaints/reports moderation, and admin-gated "
       "operational controls.", bold_lead="Operational Backbone: ")

# ============================================================ ASSUMPTIONS
heading("ASSUMPTIONS", 1)
para("The work scope, architecture and estimate below assume the current "
     "technology stack is retained (no rewrite). The remediation hardens and "
     "completes the existing system rather than replacing it. Final figures "
     "are subject to the open questions listed under Practical Limitations.")

heading("Complete Technical Stack", 3)
bullet("Expo / React Native (TypeScript), expo-router file-based navigation, "
       "Zustand + React Context state.", bold_lead="Mobile: ")
bullet("NestJS 10 (modular, controller → service → Prisma), REST + Socket.IO "
       "gateways.", bold_lead="Backend: ")
bullet("PostgreSQL via Prisma (single schema, snake_case, soft-delete via "
       "deleted_at).", bold_lead="Database: ")
bullet("Firebase ID-token verification (firebase-admin) + admin guard; "
       "ValidationPipe; planned rate-limiting + Helmet.", bold_lead="Auth & Security Layer: ")
bullet("Stripe Connect (destination charges, organizer payouts, webhooks with "
       "signature verification + idempotency).", bold_lead="Payments: ")
bullet("Railway (target); CI/CD gate (build + lint + test); Redis pub/sub for "
       "multi-instance Socket.IO + cache (growth-ready 10k+).",
       bold_lead="Infrastructure: ")
bullet("Sentry (already initialized); SendGrid for transactional email "
       "(assumption — to confirm).", bold_lead="Monitoring & Comms: ")

para("Security & Compliance:", bold=True)
bullet("at-rest + TLS in transit; secrets resolved from environment, validated "
       "at boot.", bold_lead="Encryption: ")
bullet("Firebase ID token on every authenticated request; admin routes "
       "fail-closed.", bold_lead="Authentication: ")
bullet("UK/EU GDPR in scope — data export/erasure with a 6-year financial-"
       "ledger retention carve-out (HMRC). AML handled by Stripe.",
       bold_lead="Compliance: ")
bullet("money stored as integer pence (no float-for-money), parameterized SQL, "
       "Stripe webhook idempotency verified clean.", bold_lead="Money Integrity: ")

heading("System Architecture", 3)
para("Two independent deployables share one Firebase-authenticated user "
     "identity. The mobile app talks to the backend over REST "
     "(EXPO_PUBLIC_API_URL) and Socket.IO; the backend verifies the Firebase "
     "ID token, owns all business logic and persistence, and integrates Stripe "
     "for the payment lifecycle. Hardening adds a CI/CD gate, graceful "
     "shutdown, environment validation, database indexing/soft-delete "
     "enforcement, and Redis-backed horizontal scale.")
bullet("Expo/React Native app — golfer-facing UI, the only client.",
       bold_lead="Mobile Client: ")
bullet("NestJS REST + WebSocket API — auth, games, payments, chat, "
       "notifications, moderation.", bold_lead="Backend API: ")
bullet("Stripe Connect, Firebase, Mapbox, push notifications, transactional "
       "email.", bold_lead="External Integrations: ")

# ============================================================ ROADMAP & PHASES
heading("ROADMAP & PHASES", 1)

heading("In Scope Functions", 3)
para("A. Security & Access Control", bold=True)
bullet("close IDOR on messages/conversations/complaints; guard /users CRUD, "
       "notifications/send-all and leaderboard writes; authenticate WebSocket "
       "room joins; fix reflected XSS on /go.", bold_lead="Hardening: ")
para("B. Payments & Payouts", bold=True)
bullet("one canonical payment flow; server-side Stripe verification (no "
       "client-trusted amounts); 2-day hold + automated payout; "
       "refund-on-cancel; atomic, idempotent, double-payout-safe writes.",
       bold_lead="Completion: ")
para("C. Data Layer", bold=True)
bullet("enforce soft-delete globally; add FK / hot-path indexes (concurrent "
       "migrations); fix unique constraints ignoring deleted_at; pagination on "
       "list endpoints; transactional financial writes.", bold_lead="Hardening: ")
para("D. Mobile App", bold=True)
bullet("fix API contract drift (EVENING time-slot, CONFIRMED status); 401 "
       "handling + token refresh; ErrorBoundary; remove fake Stripe key "
       "fallback; EAS env wiring.", bold_lead="Hardening: ")
para("E. Reliability & Infrastructure", bold=True)
bullet("CI/CD gate (build + lint + test); graceful shutdown + Prisma "
       "disconnect; env validation at boot; hardened Dockerfile; "
       "Helmet/throttler/CORS; Redis adapter for multi-instance scale.",
       bold_lead="Platform: ")
para("F. Compliance, Legal & Comms", bold=True)
bullet("GDPR data export/erasure with financial-ledger retention; store-"
       "mandatory privacy/consent flows; transactional email (SendGrid).",
       bold_lead="UK/EU: ")
para("G. Performance & Scale", bold=True)
bullet("remove N+1 in payout review; cache suggested/nearby; search indexing "
       "(pg_trgm); rate-limit defaults; WebSocket scale; moderation "
       "enforcement; mobile react-query.", bold_lead="Growth-ready 10k+: ")
para("H. Quality & Testing", bold=True)
bullet("coverage gate for money/auth/games critical flows; mobile smoke + "
       "critical-flow tests (built from near-zero).", bold_lead="Quality Gate: ")

heading("Roadmap", 3)
para("Agile delivery in 2-week sprints over ~8 weeks (2 months) with a "
     "dedicated parallel squad — compressed from a sequential 4-month plan by "
     "running three backend tracks plus mobile concurrently. Phase 1 (P0 "
     "blockers) gates everything; once blockers clear, the remaining "
     "workstreams run in parallel across four tracks (T1 Security & Payments, "
     "T2 Data/Performance/Compliance, T3 Infra & DevOps, T4 Mobile), with "
     "testing threaded throughout.")
def sprint(name, focus, deliverables):
    para(name, bold=True, color=GREEN)
    bullet(focus, bold_lead="Focus: ")
    para("Key Deliverables:", bold=False, italic=True)
    for d in deliverables:
        bullet(d, level=1)

sprint("Sprint 1 (Weeks 1–2) — P0 Blockers + CI Foundation",
       "eliminate every go-live-blocking exploit and stand up the quality gate "
       "before parallel work begins.",
       ["T1 Backend: close all IDOR/access-control holes, authenticate "
        "sockets, guard admin/user/leaderboard routes, fix client-trusted "
        "payment verification, make double-payout atomic/idempotent.",
        "T3 Backend: graceful shutdown + Prisma disconnect; env validation.",
        "T3 DevOps: stand up CI/CD gate (build + lint + test).",
        "T4 Mobile: begin API-contract reconciliation in parallel."])
sprint("Sprint 2 (Weeks 3–4) — Payments, Data & Infra (parallel)",
       "one canonical, atomic, automated payment lifecycle on a hardened, "
       "scalable platform.",
       ["T1 Backend: canonical flow, 2-day hold + auto-payout, refund-on-cancel.",
        "T2 Backend: soft-delete enforcement, FK/hot-path indexes, pagination.",
        "T3 Infra: hardened Dockerfile, Helmet/throttler/CORS, Redis adapter "
        "for Socket.IO + cache (multi-instance).",
        "T4 Mobile: 401/refresh, ErrorBoundary, remove fake key, EAS env."])
sprint("Sprint 3 (Weeks 5–6) — Compliance, Scale & Completeness (parallel)",
       "UK/EU + store readiness, growth-ready performance, production-clean "
       "code; test build-out ramps up.",
       ["T2 Backend: GDPR export/erasure with ledger retention; consent/privacy "
        "flows; SendGrid transactional email.",
        "T2 Backend: remove N+1, caching, pg_trgm search index, rate-limit, "
        "WebSocket scale, moderation enforcement.",
        "T1/T2 Backend: live-stub completion, timezone correctness, CVE "
        "patching, /v1/games sunset.",
        "T4 Mobile: react-query data layer + game-performance optimizations.",
        "QC: coverage build-out for money/auth/games + mobile smoke tests."])
sprint("Sprint 4 (Weeks 7–8) — Quality Gate, UAT & Go-Live",
       "verified, tested, deployed.",
       ["QC: coverage gate enforced for money/auth/games; mobile critical-flow "
        "tests green in CI.",
        "All: staging deploy, UAT, bug-fix, production deploy + handover docs."])

heading("Practical Limitations", 3)
bullet("final figures depend on confirming the canonical payment flow "
       "(/games vs /v1/games) and whether the client-trusted endpoint is "
       "still live.", bold_lead="Payment-flow decision: ")
bullet("whether public endpoints (leaderboards, round, locations, "
       "courses/profiles CRUD) are intended features or leftover scaffold.",
       bold_lead="Endpoint intent: ")
bullet("Railway deploy target and instance count (drives the Redis adapter "
       "decision); EAS production env vars must be set on the dashboard.",
       bold_lead="Deploy targets: ")
bullet("financial-data retention period and email provider (SendGrid assumed); "
       "search backend (pg_trgm assumed); rate-limit thresholds.",
       bold_lead="Assumptions to confirm: ")
bullet("mobile testing is built from near-zero (~162 files / ~34k LOC, native "
       "Stripe/Firebase/Mapbox mocks) — the single largest variable in the "
       "range.", bold_lead="Mobile test build-out: ")

heading("Out-of-Scope Features", 3)
bullet("god-service refactor (games ~2,200 LOC) — high regression risk on "
       "money code; deferred post-launch.", bold_lead="Large refactors: ")
bullet("broad/strict-mode test coverage beyond money/auth/critical flows — "
       "deferred post-launch.", bold_lead="Exhaustive coverage: ")
bullet("accessibility (a11y) and internationalisation (i18n) — deferred "
       "post-launch backlog (~11–17 md when undertaken).",
       bold_lead="a11y & i18n: ")
bullet("no new product features; engagement is hardening/completion of the "
       "existing scope only.", bold_lead="New features: ")
bullet("Elasticsearch/full-text search; native AML (handled by Stripe).",
       bold_lead="Heavy infra: ")

# ============================================================ DELIVERABLES
heading("DELIVERABLES", 1)
para("On completion the Client receives a production-ready system plus the "
     "supporting artefacts:")
for d in [
    ("Hardened backend + mobile source code", "all remediation merged to main, "
     "passing the CI/CD quality gate."),
    ("Go-live gate sign-off", "0 open Critical security/money findings; "
     "checklist evidence per phase."),
    ("Automated payment lifecycle", "single canonical flow with hold, payout, "
     "refund — atomic and idempotent."),
    ("Database migrations", "indexes, soft-delete enforcement, constraint "
     "fixes (concurrent, reversible)."),
    ("CI/CD + infrastructure config", "pipeline, hardened Dockerfile, env "
     "validation, Redis-backed scale."),
    ("Test suites & coverage report", "money/auth/games gate + mobile smoke / "
     "critical-flow tests."),
    ("Compliance package", "GDPR export/erasure, consent flows, transactional "
     "email."),
    ("Deployment & handover guide", "runbook, environment matrix, UAT report, "
     "operational docs."),
]:
    bullet(d[1], bold_lead=d[0] + ": ")

# ============================================================ SCHEDULE & COST
heading("SCHEDULE & COST", 1)

heading("Working Breakdown Structure", 2)
para("Delivery is agile (2-week sprints) with a parallel squad. The WBS below "
     "groups the remediation scope by workstream; effort is point-estimated per "
     "item and rolled into the cost model.")

wbs_headers = ["#", "Category", "Function", "Sub-Function", "Notes"]
wbs_rows = [
    ["A. Security & Access Control"],
    ["A1", "Access control", "IDOR remediation", "Messages, conversations, complaints (PII)"],
    ["A2", "Access control", "Route guards", "/users CRUD, notifications/send-all, leaderboard writes"],
    ["A3", "Real-time", "Socket authentication", "Authenticate room joins (no eavesdrop/inject)"],
    ["A4", "Web", "Reflected XSS fix", "/go attribution endpoint"],
    ["B. Payments & Payouts"],
    ["B1", "Payments", "Server-side verification", "No client-trusted amount/intent"],
    ["B2", "Payments", "Canonical flow", "Resolve /games vs /v1/games"],
    ["B3", "Payments", "Hold + auto-payout", "2-day hold, scheduled payout, refund-on-cancel"],
    ["B4", "Payments", "Atomicity & idempotency", "Double-payout-safe, transactional writes"],
    ["C. Data Layer"],
    ["C1", "Database", "Soft-delete enforcement", "Global; honour deleted_at on reads"],
    ["C2", "Database", "Indexing", "FK + hot-path indexes, concurrent migrations"],
    ["C3", "Database", "Constraints + pagination", "Unique-with-deleted_at; list pagination"],
    ["D. Mobile App"],
    ["D1", "Mobile", "Contract reconciliation", "EVENING time-slot, CONFIRMED status drift"],
    ["D2", "Mobile", "Auth resilience", "401 handling + token refresh"],
    ["D3", "Mobile", "Stability", "ErrorBoundary; remove fake Stripe key; EAS env"],
    ["E. Reliability & Infrastructure"],
    ["E1", "DevOps", "CI/CD gate", "Build + lint + test blocks deploy"],
    ["E2", "Infra", "Runtime hardening", "Graceful shutdown, env validation, Dockerfile"],
    ["E3", "Infra", "Edge hardening", "Helmet, throttler, CORS, Swagger gating"],
    ["E4", "Infra", "Horizontal scale", "Redis adapter for Socket.IO + cache"],
    ["F. Compliance, Legal & Comms"],
    ["F1", "Compliance", "GDPR export/erasure", "Ledger retention carve-out (6 yr)"],
    ["F2", "Compliance", "Consent & privacy flows", "Store-mandatory readiness"],
    ["F3", "Comms", "Transactional email", "SendGrid integration"],
    ["G. Performance & Scale"],
    ["G1", "Performance", "Query optimization", "Remove N+1, caching"],
    ["G2", "Scale", "Search + rate-limit", "pg_trgm index, rate-limit defaults"],
    ["G3", "Scale", "WS scale + moderation", "WebSocket scale, moderation enforcement"],
    ["G4", "Mobile", "react-query + game perf", "Client data layer + perf"],
    ["H. Quality & Testing"],
    ["H1", "QA", "Coverage gate", "Money/auth/games critical flows"],
    ["H2", "QA", "Mobile tests", "Smoke + critical-flow (from near-zero)"],
    ["I. Out of scope (deferred post-launch)"],
    ["I1", "Refactor", "God-service refactor", "<Not in scope>"],
    ["I2", "QA", "Broad/strict coverage", "<Not in scope>"],
    ["I3", "UX", "a11y + i18n", "<Not in scope>"],
]
# compute which indices are section bands
section_idx = {i for i, r in enumerate(wbs_rows) if len(r) == 1}
add_table(wbs_headers, wbs_rows, widths=[0.4, 1.5, 1.7, 2.8, 0.0],
          section_rows=section_idx, font_size=9)

heading("Cost & Budget estimation", 2)
para("Investment required to take Alba from its current state to production "
     "go-live, delivered by a dedicated SotaTek squad. Pricing is time-and-"
     "materials by role in man-months (M/M); USD.")
para("Recommended Scope — Production Go-Live (~8 weeks / 2 months)", bold=True,
     color=GREEN)
para("Customer-confirmed scope = full hardening (Phases 1–11, growth-ready "
     "10k+, UK/EU compliance, production-grade stubs, Redis) excluding the "
     "deferred god-service refactor and broad/strict test coverage. Total "
     "engineering effort ≈ 150–186 man-days; the table prices the point "
     "estimate within that range. Calendar time is halved versus a sequential "
     "plan by running three backend tracks plus mobile in parallel — the total "
     "effort and cost are unchanged.", italic=True, size=10)

cost_headers = ["Position", "Month 1", "Month 2",
                "Total (M/M)", "Unit Price (USD)", "Total Cost (USD)"]
# (role, [m1, m2], unit_price)
cost_data = [
    ("Project Manager",          [0.5, 0.5], 3500),
    ("BA & UI/UX Designer",      [0.5, 0.25], 3000),
    ("Backend Engineer (x3)",    [3.0, 3.0], 3000),
    ("Mobile Engineer",          [0.75, 0.75], 3000),
    ("QC Engineer",              [0.5, 1.0], 2800),
    ("DevOps Engineer",          [0.5, 0.25], 3600),
]
cost_rows = []
tot_mm = 0.0
tot_cost = 0.0
month_tot = [0.0, 0.0]
for role, months, unit in cost_data:
    mm = round(sum(months), 2)
    cost = mm * unit
    tot_mm += mm
    tot_cost += cost
    for k in range(2):
        month_tot[k] += months[k]
    cost_rows.append([role] + [f"{m:g}" for m in months] +
                     [f"{mm:g}", f"{unit:,}", f"{cost:,.0f}"])
cost_rows.append(
    ["TOTAL"] + [f"{m:g}" for m in month_tot] +
    [f"{tot_mm:g}", "", f"{tot_cost:,.0f}"]
)
add_table(cost_headers, cost_rows,
          widths=[1.9, 0.7, 0.7, 0.9, 1.1, 1.2], font_size=9)
para()
para(f"Recommended-scope investment: USD {tot_cost:,.0f}  "
     f"(≈ {tot_mm:g} man-months across the squad, delivered in ~2 months).",
     bold=True)
bullet("the deferred god-service refactor and broad/strict coverage are "
       "available as a post-launch add-on (~14–18 md, ~USD 6,000–7,500) once "
       "the product is live.", bold_lead="Optional add-on: ")
bullet("a11y + i18n localisation can be scoped separately post-launch "
       "(~11–17 md).", bold_lead="Future phase: ")
bullet("a ~6-week (1.5-month) variant is feasible if compliance (F) and "
       "growth-scale (G) are deferred to a fast-follow and the launch targets "
       "the go-live blockers only (A–E + targeted H).",
       bold_lead="Aggressive timeline: ")
para("Estimate is a presale range; final figures confirmed after the open "
     "questions above are answered (±~5 md).", italic=True, size=10)

# ============================================================ PAYMENT MILESTONE
heading("PAYMENT MILESTONE", 1)
para("Payments are tied to verifiable delivery milestones; total equals 100% "
     "of the engagement fee.")
ms_headers = ["Milestone", "Deliverable", "Percentage", "Due Timeline"]
ms_rows = [
    ["M1: Contract Signing", "Signed agreement, squad mobilisation, kickoff",
     "40%", "Upon contract execution"],
    ["M2: P0 Blockers Demo",
     "0 open Critical findings; payments + access-control demo on dev",
     "30%", "End of Sprint 2 (~Week 4)"],
    ["M3: Staging & UAT", "Full scope on staging; UAT begins",
     "20%", "End of Sprint 3 (~Week 6)"],
    ["M4: Go-Live", "UAT sign-off, production deploy, handover docs",
     "10%", "End of Sprint 4 (~Week 8)"],
]
add_table(ms_headers, ms_rows, widths=[1.6, 3.0, 0.9, 1.7], font_size=9.5)
para()
para("Total: 100%", bold=True)

doc.save(OUTPUT)
print("Saved", OUTPUT, "| total cost USD", f"{tot_cost:,.0f}", "| MM", tot_mm)
